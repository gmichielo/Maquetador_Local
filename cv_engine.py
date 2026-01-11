import pdfplumber
from docx import Document
from docx2pdf import convert
import platform
import re
import unicodedata
import shutil
import os
import time
import threading
import pythoncom

# =========================================================
# 1. UTILIDADES
# =========================================================

def normalize_text(text):
    text = unicodedata.normalize("NFKD", text)
    text = "".join(c for c in text if not unicodedata.combining(c))
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{2,}', '\n', text)
    return text.strip()


def read_pdf(path):
    blocks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text(layout=True)
            if txt:
                blocks.append(txt)
    return "\n".join(blocks)


# =========================================================
# 2. NORMALIZACIÓN ATS
# =========================================================

def rebuild_structure(text):
    headers = [
        "perfil profesional", "profile",
        "experiencia laboral", "work experience",
        "education", "educacion",
        "skills", "habilidades",
        "languages", "idiomas"
    ]

    for h in headers:
        text = re.sub(rf"\s*{h}\s*", f"\n{h.upper()}\n", text, flags=re.IGNORECASE)

    # Insertar salto de línea antes de cualquier • (u otros bullets) y mantener el símbolo
    text = re.sub(r"\s*([•\*\|▪●])\s*", r"\n\1 ", text)
    
    return normalize_text(text)

def split_lines(text):
    return [l.strip() for l in text.split("\n") if len(l.strip()) > 2]

def normalize_experience_lines(lines):
    cleaned = []
    for l in lines:
        l = l.replace("|", "").strip()
        l = re.sub(r'^[●•\-]\s*', '', l)
        cleaned.append(l)
    return cleaned

# =========================================================
# 3. EXTRACCIÓN
# =========================================================

def extract_name(lines):
    """
    Extrae el nombre de un CV, intentando ser más flexible.
    """
    for line in lines[:10]:  # solo revisar primeras 10 líneas
        line_clean = line.strip()
        # ignorar líneas vacías o con números o emails
        if not line_clean:
            continue
        if re.search(r'\d', line_clean):
            continue
        if re.search(r'\S+@\S+', line_clean):  # ignorar emails
            continue
        if len(line_clean.split()) >= 2:  # debe tener al menos 2 palabras
            # eliminar títulos o palabras comunes como "perfil", "experiencia"
            if re.search(r'perfil|experiencia|skills|educacion', line_clean, re.IGNORECASE):
                continue
            return line_clean  # devuelve tal cual
    return "Nombre no detectado"


def extract_contact(text):
    email = re.search(r'\S+@\S+', text)

    phone = re.search(
        r'(\+\d{1,3}[\s\-]?)?(\(?\d{2,4}\)?[\s\-]?)?\d{3,4}[\s\-]?\d{3,4}',
        text
    )

    github = re.search(r'github\.com/\S+', text, re.IGNORECASE)
    linkedin = re.search(r'https?://(www\.)?linkedin\.com/\S+', text, re.IGNORECASE)

    telefono = phone.group(0).replace("\n", " ").strip() if phone else ""

    return {
        "email": email.group(0) if email else "",
        "telefono": telefono,
        "github": github.group(0) if github else "",
        "linkedin": linkedin.group(0) if linkedin else ""
    }


def extract_certificaciones(educacion_lines):
    certs = []
    edu = []

    for l in educacion_lines:
        low = l.lower()
        if any(k in low for k in ["cert", "ibm", "caelum", "oracle", "aws"]):
            certs.append(l)
        else:
            edu.append(l)

    return edu, certs


SECTIONS = {
    "perfil": ["perfil profesional", "profile", "summary", "about me"],
    "experiencia": ["experiencia laboral", "work experience", "experiencia profesional", "work history"],
    "educacion": ["educacion", "education", "certificaciones", "formacion academica", "education and training"],
    "skills": ["skills", "habilidades", "experticia tecnica", "competencias", "skills & competencies"],
    "idiomas": ["idiomas", "languages", "language","language skills"],
    "proyectos": ["proyectos", "proyectos destacados", "projects"]
}


def split_by_sections(lines):
    data = {k: [] for k in SECTIONS}
    current = None

    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue

        low = line_clean.lower()
        matched = False

        for k, keys in SECTIONS.items():
            for key in keys:
                # 👇 HEADER si la línea EMPIEZA por el nombre del header
                if low.startswith(key):
                    current = k
                    matched = True
                    break
            if matched:
                break

        if not matched and current:
            data[current].append(line_clean)

    return data



def extract_skills(lines):
    skills = []
    seen = set()

    for l in lines:
        for s in re.split(r"[,:]", l):
            s = s.strip()

            if len(s) <= 2:
                continue

            # evitar duplicados manteniendo orden
            key = s.lower()
            if key in seen:
                continue

            seen.add(key)
            skills.append(s)

    return skills


def extract_idiomas(lines):
    idiomas = {}

    for i, l in enumerate(lines):
        low = l.lower()

        # ---- FORMATO NORMAL (Español: Nativo) ----
        matches = re.findall(r'([A-Za-zÁÉÍÓÚáéíóú]+)\s*[:\-]\s*(\w+)', l)
        for lang, lvl in matches:
            idiomas[lang.capitalize()] = lvl.capitalize()

        # ---- EUROPASS: Mother tongue(s) ----
        if "mother tongue" in low:
            if i + 1 < len(lines):
                idiomas[lines[i + 1].strip().capitalize()] = "Nativo"

        # ---- EUROPASS: ENGLISH C1 C1 C1 ----
        m = re.match(r'^([A-Z]+)\s+(A1|A2|B1|B2|C1|C2)', l)
        if m:
            idiomas[m.group(1).capitalize()] = m.group(2)

    return idiomas

DATE_REGEX = re.compile(
    r"""
    (
        # 03/2025 - 09/2025
        \d{2}/\d{4}\s*[-–]\s*(\d{2}/\d{4}|actualidad|present|current)
        |
        # 2013 - 2014
        \d{4}\s*[-–]\s*(\d{4}|actualidad|present|current)
        |
        # Mar 2015 - Sep 2017
        (ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic|
         enero|febrero|marzo|abril|mayo|junio|julio|agosto|
         septiembre|setiembre|octubre|noviembre|diciembre|
         jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)
        \s+\d{4}\s*[-–]\s*
        (
            (ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic|
             enero|febrero|marzo|abril|mayo|junio|julio|agosto|
             septiembre|setiembre|octubre|noviembre|diciembre|
             jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)
            \s+\d{4}
            |
            actualidad|present|current
        )
    )
    """,
    re.IGNORECASE | re.VERBOSE
)

def format_experiencia_bloques(bloques):
    salida = []
    for b in bloques:
        salida.append(
            f"""{b['fecha']}
Empresa: {b['empresa']}
Puesto: {b['puesto']}
Funciones:
""" + "\n".join(f"• {f}" for f in b["funciones"])
        )
    return "\n\n".join(salida)

def format_experiencia_plantilla(lines):
    bloques = []
    actual = None

    for l in lines:
        l = l.strip()
        if not l:
            continue

        # -------- FECHA SOLA --------
        if DATE_REGEX.search(l) and len(l.split()) <= 6:
            if actual:
                actual["fecha"] = l
            continue

        # -------- PUESTO + FECHA (EUROPASS) --------
        if DATE_REGEX.search(l) and " – " in l:
            if actual:
                actual["puesto"] = l.split(" – ")[0].strip()
                actual["fecha"] = " – ".join(l.split(" – ")[1:])
            continue

        # -------- EMPRESA (EUROPASS ICONO / TEXTO) --------
        if "–" in l and any(city in l.lower() for city in ["madrid", "gijon", "oviedo", "spain"]):
            if actual:
                bloques.append(actual)
            actual = {
                "empresa": l.replace("", "").strip(),
                "puesto": "",
                "fecha": "",
                "funciones": []
            }
            continue

        # -------- EMPRESA CLÁSICA --------
        if l.isupper() and len(l.split()) <= 6:
            if actual:
                bloques.append(actual)
            actual = {
                "empresa": l,
                "puesto": "",
                "fecha": "",
                "funciones": []
            }
            continue

        if not actual:
            continue

        # -------- PUESTO --------
        if not actual["puesto"]:
            actual["puesto"] = l
            continue

        # -------- FUNCIONES --------
        actual["funciones"].append(l)

    if actual:
        bloques.append(actual)

    # -------- FORMATO FINAL --------
    salida = []
    for b in bloques:
        salida.append(
            f"""{b['fecha']}
Empresa: {b['empresa']}
Puesto: {b['puesto']}
Funciones:
""" + "\n".join(f"• {f}" for f in b["funciones"])
        )

    return "\n\n".join(salida)

def format_proyectos(lines):
    bloques = []
    actual = []

    for l in lines:
        # Título del proyecto (línea corta o sin bullets)
        if not l.startswith(("•", "-", "*")) and len(l.split()) <= 6:
            if actual:
                bloques.append("\n".join(actual))
                actual = []
            actual.append(l)
        else:
            actual.append(l)

    if actual:
        bloques.append("\n".join(actual))

    return "\n\n".join(bloques)

def clean_bullets(lines):
    return [re.sub(r'^[•\-\*]\s*', '', l) for l in lines]

# =========================================================
# 4. PARSER PRINCIPAL
# =========================================================

def is_europass(text):
    markers = [
        "europass",
        "mother tongue",
        "language skills",
        "education and training"
    ]
    low = text.lower()
    return any(m in low for m in markers)

def parse_experiencia_europass(lines):
    bloques = []

    empresa = None
    puesto = None
    fecha = None
    funciones = []

    for l in lines:
        l = l.strip()
        if not l:
            continue

        # Empresa (empresa – ciudad, país)
        if " – " in l and any(x in l.lower() for x in ["spain", "madrid", "gijon", "oviedo"]):
            if empresa:
                bloques.append({
                    "empresa": empresa,
                    "puesto": puesto or "",
                    "fecha": fecha or "",
                    "funciones": funciones
                })
            empresa = l.replace("", "").strip()
            puesto = None
            fecha = None
            funciones = []
            continue

        # Puesto + fecha
        if DATE_REGEX.search(l) and " – " in l:
            parts = l.split(" – ")
            puesto = parts[0].strip()
            fecha = " – ".join(parts[1:])
            continue

        # Funciones
        if l.startswith(("•", "-", "*")) or len(l.split()) > 4:
            funciones.append(l.lstrip("•-* ").strip())

    if empresa:
        bloques.append({
            "empresa": empresa,
            "puesto": puesto or "",
            "fecha": fecha or "",
            "funciones": funciones
        })

    return bloques

def parse_cv(pdf_path):
    raw = read_pdf(pdf_path)
    structured = rebuild_structure(raw)
    lines = split_lines(structured)
    sections = split_by_sections(lines)

    if is_europass(raw):
        bloques = parse_experiencia_europass(sections["experiencia"])
        experiencia_formateada = format_experiencia_bloques(bloques)
        experiencia = sections["experiencia"]
    else:
        sections["experiencia"] = normalize_experience_lines(sections["experiencia"])
        experiencia = sections["experiencia"]
        experiencia_formateada = format_experiencia_plantilla(experiencia)

    educacion_limpia, certificaciones = extract_certificaciones(sections["educacion"])

    return {
        "nombre": extract_name(lines),
        "contacto": extract_contact(structured),
        "perfil": " ".join(sections["perfil"]),
        "skills": extract_skills(sections["skills"]),
        "experiencia": experiencia,
        "experiencia_formateada": experiencia_formateada,
        "educacion": educacion_limpia,
        "certificaciones": certificaciones,
        "idiomas": extract_idiomas(sections["idiomas"]),
        "proyectos": sections["proyectos"],
        "proyectos_formateados": format_proyectos(sections["proyectos"])
    }


# =========================================================
# 5. DOCX / PDF
# =========================================================

def cv_json_to_docx_data(cv):
    return {
        "NOMBRE": cv["nombre"],
        "EMAIL": cv["contacto"]["email"],
        "TELEFONO": cv["contacto"]["telefono"],
        "GITHUB": cv["contacto"]["github"],
        "LINKEDIN": cv["contacto"]["linkedin"],
        "PERFIL": cv["perfil"],
        "SKILLS": ", ".join(cv["skills"]),
        "FORMACION": "\n".join(cv["educacion"]),
        "EDUCACION": "\n".join(cv["educacion"]),
        "CERTIFICACIONES": "\n".join(cv["certificaciones"]),
        "EXPERIENCIA": cv["experiencia"],
        "EXPERIENCIA_PLANTILLA": cv["experiencia_formateada"],
        "IDIOMAS": "\n".join(f"• {k}: {v}" for k, v in cv["idiomas"].items()),
        "PROYECTOS": cv.get("proyectos_formateados", "")
    }

def is_empty_value(v):
    if v is None:
        return True
    if isinstance(v, str) and not v.strip():
        return True
    if isinstance(v, (list, dict)) and len(v) == 0:
        return True
    return False

def replace_placeholders(doc, data, empty_text=""):
    # --------- PÁRRAFOS ---------
    for p in doc.paragraphs:
        full_text = p.text

        for k, v in data.items():
            placeholder = f"{{{{{k}}}}}"

            if placeholder not in full_text:
                continue

            if is_empty_value(v):
                # Si el párrafo SOLO contiene el placeholder → borrar párrafo
                if full_text.strip() == placeholder:
                    p.text = ""
                else:
                    # Si está mezclado con texto → reemplazar por texto alternativo
                    p.text = full_text.replace(placeholder, empty_text)
            else:
                p.text = full_text.replace(placeholder, str(v))

    # --------- TABLAS ---------
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    full_text = p.text

                    for k, v in data.items():
                        placeholder = f"{{{{{k}}}}}"

                        if placeholder not in full_text:
                            continue

                        if is_empty_value(v):
                            if full_text.strip() == placeholder:
                                p.text = ""
                            else:
                                p.text = full_text.replace(placeholder, empty_text)
                        else:
                            p.text = full_text.replace(placeholder, str(v))

def replace_placeholders_preserve_style(doc, data, empty_text=""):
    def replace_in_runs(runs, data):
        for run in runs:
            for k, v in data.items():
                placeholder = f"{{{{{k}}}}}"

                if placeholder not in run.text:
                    continue

                if is_empty_value(v):
                    run.text = run.text.replace(placeholder, empty_text)
                else:
                    run.text = run.text.replace(placeholder, str(v))

    # --------- PÁRRAFOS ---------
    for p in doc.paragraphs:
        replace_in_runs(p.runs, data)

    # --------- TABLAS ---------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_runs(p.runs, data)


def generate_cv_from_template(template_path, cv_json, output_dir="output"):
    """
    Genera un DOCX y un PDF desde la plantilla usando docx2pdf.
    Usa threading + pythoncom.CoInitialize() para que funcione en Flask.
    """
    os.makedirs(output_dir, exist_ok=True)

    # -----------------------------
    # Preparar nombres únicos
    # -----------------------------
    safe_name = cv_json["nombre"].replace(" ", "_") or "CV"
    timestamp = int(time.time())
    docx_out = os.path.abspath(os.path.join(output_dir, f"CV_{safe_name}_{timestamp}.docx"))
    pdf_out = os.path.abspath(os.path.join(output_dir, f"CV_{safe_name}_{timestamp}.pdf"))
    template_path = os.path.abspath(template_path)

    # -----------------------------
    # Limpiar archivos antiguos (opcional)
    # -----------------------------
    if os.path.exists(docx_out):
        os.remove(docx_out)
    if os.path.exists(pdf_out):
        os.remove(pdf_out)

    # -----------------------------
    # Copiar plantilla y reemplazar placeholders
    # -----------------------------
    shutil.copy(template_path, docx_out)
    doc = Document(docx_out)

    data = cv_json_to_docx_data(cv_json)

    # Reemplazo de placeholders
    replace_placeholders_preserve_style(doc, data)

    doc.save(docx_out)

    # -----------------------------
    # Función para generar PDF en hilo separado
    # -----------------------------
    pdf_generated = False

    def convert_pdf_thread():
        nonlocal pdf_generated
        try:
            pythoncom.CoInitialize()  # Inicializar COM en este hilo
            convert(docx_out, pdf_out)
            if os.path.exists(pdf_out):
                pdf_generated = True
        except Exception as e:
            print("Error generando PDF en hilo:", e)

    if platform.system().lower() == "windows":
        thread = threading.Thread(target=convert_pdf_thread)
        thread.start()
        thread.join()  # Esperamos a que termine el PDF

        if not pdf_generated:
            print("PDF no se pudo generar después de intentar en hilo.")
            pdf_out = None
    else:
        # En otros sistemas no se usa docx2pdf
        pdf_out = None

    # -----------------------------
    # Devolver DOCX siempre, PDF si se generó
    # -----------------------------
    return docx_out, pdf_out if pdf_generated else None